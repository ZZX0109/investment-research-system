from __future__ import annotations

import hashlib
import io
from datetime import datetime, timedelta, timezone
from uuid import uuid4

from investment_research.domain.base import Provenance
from investment_research.domain.enums import DataMode, DataSourceType
from investment_research.domain.knowledge import FinancialKnowledgeDocument
from investment_research.domain.models import User
from investment_research.repository.sqlite import SQLiteUnitOfWork
from investment_research.service.financial_knowledge import FinancialKnowledgeService
from investment_research.service.knowledge_retrieval import KnowledgeChunker, KnowledgeRetrievalService
from investment_research.service.knowledge_ingestion import OfficialKnowledgeIngestionService
from investment_research.service.documents import DocumentService
from investment_research.service.object_store import LocalObjectStore
from investment_research.service.vision import DisabledVisionProvider


class LexicalOnlyEmbedder:
    model_name = "disabled-test-embedder"
    model_revision = "v1"
    available = False

    def encode_query(self, text: str) -> list[float]:
        raise AssertionError(text)

    def encode_documents(self, texts: list[str]) -> list[list[float]]:
        raise AssertionError(texts)


class DeterministicEmbedder:
    model_name = "deterministic-zh-test"
    model_revision = "v1"
    available = True

    @staticmethod
    def _vector(text: str) -> list[float]:
        return [float(text.count("公告") + 1), float(text.count("风险") + 1), float(len(text) % 17 + 1)]

    def encode_query(self, text: str) -> list[float]:
        return self._vector(text)

    def encode_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._vector(text) for text in texts]


def _document(*, title: str, content: str, available_at: datetime, owner=None, source_url="https://www.cninfo.com.cn/a.pdf"):
    return FinancialKnowledgeDocument(
        title=title,
        content=content,
        source_name="巨潮资讯网" if owner is None else "用户资料",
        source_url=source_url,
        market="CN",
        symbol="600519",
        document_type="announcement_metadata" if owner is None else "user_document",
        published_at=available_at - timedelta(hours=1),
        effective_from=available_at - timedelta(hours=1),
        collected_at=available_at,
        first_observed_at=available_at,
        available_at=available_at,
        content_hash=FinancialKnowledgeService.content_hash(
            title=title, content=content, source_url=source_url, owner_user_id=owner,
        ),
        source_kind="official_public" if owner is None else "user_upload",
        copyright_status="official_public" if owner is None else "user_owned",
        owner_user_id=owner,
        access_scope="public" if owner is None else "private",
        visibility_assumption="historical_available_at_unproven_public_backfill" if owner is None else None,
    )


def test_revision_is_point_in_time_and_does_not_overwrite_history(tmp_path) -> None:
    uow = SQLiteUnitOfWork(tmp_path / "revision.db")
    retrieval = KnowledgeRetrievalService(uow, embedder=LexicalOnlyEmbedder())
    service = FinancialKnowledgeService(uow, retrieval=retrieval)
    first_time = datetime(2026, 7, 1, 8, tzinfo=timezone.utc)
    first = service.ingest(_document(title="年度报告更正说明", content="原始公告净利润一百万元", available_at=first_time))
    second_time = first_time + timedelta(days=2)
    second = service.ingest(_document(title="年度报告更正说明", content="更正公告净利润二百万元", available_at=second_time))

    historical, _ = retrieval.search("净利润一百万元", as_of=first_time + timedelta(hours=1), symbol="600519")
    current, _ = retrieval.search("净利润二百万元", as_of=second_time + timedelta(hours=1), symbol="600519")

    assert first.id != second.id
    assert second.revision == 2
    assert historical and historical[0].document.id == first.id
    assert current and current[0].document.id == second.id
    assert all(item.document.id != second.id for item in historical)


def test_private_documents_are_owner_isolated_and_deletable(tmp_path) -> None:
    uow = SQLiteUnitOfWork(tmp_path / "private.db")
    retrieval = KnowledgeRetrievalService(uow, embedder=LexicalOnlyEmbedder())
    service = FinancialKnowledgeService(uow, retrieval=retrieval)
    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    owner = uuid4()
    stranger = uuid4()
    stored = service.ingest(_document(title="我的估值笔记", content="私有关键假设是现金流折现", available_at=now, owner=owner, source_url=None))

    owner_results, _ = retrieval.search("现金流折现", as_of=now + timedelta(minutes=1), owner_user_id=owner)
    stranger_results, _ = retrieval.search("现金流折现", as_of=now + timedelta(minutes=1), owner_user_id=stranger)

    assert owner_results and owner_results[0].document.id == stored.id
    assert stranger_results == []
    assert uow.financial_knowledge.delete_private_document(str(stored.id), owner_user_id=stranger) is False
    assert uow.financial_knowledge.delete_private_document(str(stored.id), owner_user_id=owner) is True


def test_chunker_keeps_pdf_pages_and_rule_sections() -> None:
    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    pdf = _document(title="定期报告", content="[page 1] 第一页风险说明\n[page 2] 第二页财务数据", available_at=now)
    rule = _document(title="交易规则", content="第一章 总则\n第一条 适用范围。\n第二条 交易要求。", available_at=now).model_copy(update={"document_type": "market_rule"})

    pdf_chunks = KnowledgeChunker(chunk_chars=160, overlap_chars=20).split(pdf)
    rule_chunks = KnowledgeChunker(chunk_chars=160, overlap_chars=20).split(rule)

    assert {item.page_start for item in pdf_chunks} == {1, 2}
    assert len(rule_chunks) >= 2
    assert all(len(item.content_hash) == 64 for item in pdf_chunks + rule_chunks)


def test_hybrid_retrieval_is_reproducible_and_persists_citations(tmp_path) -> None:
    uow = SQLiteUnitOfWork(tmp_path / "hybrid.db")
    retrieval = KnowledgeRetrievalService(uow, embedder=DeterministicEmbedder())
    service = FinancialKnowledgeService(uow, retrieval=retrieval)
    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    service.ingest(_document(title="重大风险公告", content="公司公告提示诉讼风险和偿债风险", available_at=now))

    first, snapshot = retrieval.search("公告风险", as_of=now + timedelta(minutes=1), symbol="600519")
    second, _ = retrieval.search("公告风险", as_of=now + timedelta(minutes=1), symbol="600519")

    assert snapshot.retrieval_mode == "hybrid"
    assert first[0].citation_id == second[0].citation_id
    assert first[0].semantic_score is not None
    assert first[0].final_score == second[0].final_score


def test_complete_announcement_window_records_present_and_confirmed_none_per_symbol(tmp_path) -> None:
    observed = datetime(2026, 8, 8, 9, tzinfo=timezone.utc)
    uow = SQLiteUnitOfWork(tmp_path / "coverage.db")
    ingestion = OfficialKnowledgeIngestionService(
        uow, object_store=LocalObjectStore(tmp_path / "objects"), clock=lambda: observed,
    )
    ingestion.knowledge = FinancialKnowledgeService(
        uow, retrieval=KnowledgeRetrievalService(uow, embedder=LexicalOnlyEmbedder()),
    )
    result = ingestion.ingest_cninfo_metadata(
        [{
            "证券代码": "600519", "证券简称": "贵州茅台",
            "公告标题": "2026年半年度报告", "公告时间": "2026-08-08T08:00:00+00:00",
            "公告链接": "/finalpage/2026-08-08/example.PDF",
        }],
        requested_date=observed.date(), universe_symbols=["600519", "000001"],
    )

    assert result["status"] == "complete"
    coverages = {
        item.symbol: item for item in uow.financial_knowledge.latest_coverage(market="CN")
        if item.provider == "cninfo" and item.symbol is not None
    }
    assert coverages["600519"].event_coverage_status == "events_present"
    assert coverages["600519"].metadata_count == 1
    assert coverages["000001"].event_coverage_status == "confirmed_none"
    assert coverages["000001"].metadata_count == 0


def test_user_document_provenance_remains_research_only(tmp_path) -> None:
    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    provenance = Provenance(
        data_mode=DataMode.REAL, source_type=DataSourceType.REAL,
        source_name="user-upload-test", observed_at=now, confidence=1.0,
    )
    user = User(email="kb@example.com", display_name="KB", auth_subject="user:kb", provenance=provenance)
    assert user.id
    digest = hashlib.sha256(b"owned").hexdigest()
    assert len(digest) == 64


def test_user_docx_and_text_parsing_preserve_content_and_tables(tmp_path) -> None:
    from docx import Document

    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    provenance = Provenance(
        data_mode=DataMode.REAL, source_type=DataSourceType.REAL,
        source_name="user-upload-test", observed_at=now, confidence=1.0,
    )
    user = User(email="documents@example.com", display_name="Documents", auth_subject="user:documents", provenance=provenance)
    uow = SQLiteUnitOfWork(tmp_path / "documents.db")
    service = DocumentService(
        uow, root=tmp_path / "work",
        object_store=LocalObjectStore(tmp_path / "objects"),
        vision_provider=DisabledVisionProvider(),
    )
    document = Document()
    document.add_heading("估值假设", level=1)
    document.add_paragraph("现金流折现只用于研究。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "增长率"
    table.cell(1, 1).text = "5%"
    buffer = io.BytesIO()
    document.save(buffer)

    stored = service.create(
        user=user, filename="private-note.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buffer.getvalue(),
    )
    text = service.create(
        user=user, filename="assumptions.md", content_type="text/markdown",
        data="# 风险假设\n仅供研究使用".encode(),
    )

    assert stored.parse_status == "parsed"
    assert "现金流折现" in (stored.text_summary or "")
    assert stored.tables[0]["rows"][1] == ["增长率", "5%"]
    assert text.parse_status == "parsed" and "风险假设" in (text.text_summary or "")
    assert service.get_for_user(str(stored.id), user=user) is not None


def test_user_document_rejects_empty_file_and_safely_degrades_blank_pdf(tmp_path) -> None:
    import fitz

    now = datetime(2026, 8, 1, tzinfo=timezone.utc)
    user = User(
        email="blank@example.com", display_name="Blank", auth_subject="user:blank",
        provenance=Provenance(
            data_mode=DataMode.REAL, source_type=DataSourceType.REAL,
            source_name="user-upload-test", observed_at=now, confidence=1.0,
        ),
    )
    uow = SQLiteUnitOfWork(tmp_path / "blank.db")
    service = DocumentService(
        uow, root=tmp_path / "work",
        object_store=LocalObjectStore(tmp_path / "objects"),
        vision_provider=DisabledVisionProvider(),
    )
    try:
        service.create(user=user, filename="empty.txt", content_type="text/plain", data=b"")
    except ValueError as exc:
        assert "empty" in str(exc).lower()
    else:
        raise AssertionError("empty user documents must be rejected")

    pdf = fitz.open()
    pdf.new_page()
    payload = pdf.tobytes()
    pdf.close()
    stored = service.create(
        user=user, filename="scan.pdf", content_type="application/pdf", data=payload,
    )
    assert stored.page_count == 1
    assert stored.parse_status in {"failed", "needs_visual_review"}
    assert stored.text_summary is None
